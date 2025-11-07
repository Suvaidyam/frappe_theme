import os
import re
from datetime import datetime

import frappe
from click.exceptions import Exit
from frappe.installer import update_site_config as _update_site_config
from frappe.utils.pdf import get_pdf


@frappe.whitelist()
def get_wf_state_by_closure(doctype, closure_type="Positive"):
	sql = """SELECT
            wfs.state,
            wfs.custom_closure
            FROM
            `tabWorkflow` AS wf
            INNER JOIN `tabWorkflow Document State` AS wfs ON wf.name = wfs.parent
            WHERE
            wf.document_type = %s AND wf.is_active = 1 AND wfs.custom_closure = %s
        """
	list = frappe.db.sql(sql, (doctype, closure_type), as_dict=1)
	if len(list) > 0:
		return list[0].state
	return None


@frappe.whitelist()
def update_site_config(key, value):
	"""
	Update a key in the site config with the given value.
	"""
	if frappe.session.user != "Administrator":
		frappe.throw("You are not authorized to perform this action.", frappe.PermissionError)

	try:
		_update_site_config(key, value)
		return frappe.get_site_config().get(key, None)
	except Exception as e:
		frappe.throw(f"Failed to update site config: {str(e)}", frappe.ValidationError)


import tempfile
from io import BytesIO

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from frappe.utils import formatdate, nowdate
from pdf2docx import Converter


def pdf_bytes_to_docx_bytes(pdf_bytes: bytes) -> bytes:
	# Convert PDF → DOCX using temp files
	with tempfile.NamedTemporaryFile(suffix=".pdf") as temp_pdf:
		temp_pdf.write(pdf_bytes)
		temp_pdf.flush()

		with tempfile.NamedTemporaryFile(suffix=".docx") as temp_docx:
			cv = Converter(temp_pdf.name)
			cv.convert(temp_docx.name)
			cv.close()

			temp_docx.seek(0)
			docx_bytes = temp_docx.read()

	# Post-process DOCX: add table borders and apply formatting
	docx_stream = BytesIO(docx_bytes)
	doc = Document(docx_stream)

	# Static formatting values for DOCX
	font_name = "Calibri"
	font_size = 12
	page_width_inch = 11.0
	page_height_inch = 8.5
	left_margin_inch = 0.5
	right_margin_inch = 0.5
	top_margin_inch = 0.7
	bottom_margin_inch = 0.7

	# Calculate available width for table
	available_width = page_width_inch - left_margin_inch - right_margin_inch
	available_width_twips = int(available_width * 1440)  # Convert inches to twips (1 inch = 1440 twips)

	# Set page size and margins for all sections
	for section in doc.sections:
		# Set page dimensions for landscape
		section.page_width = Inches(page_width_inch)
		section.page_height = Inches(page_height_inch)

		# Margins
		section.left_margin = Inches(left_margin_inch)
		section.right_margin = Inches(right_margin_inch)
		section.top_margin = Inches(top_margin_inch)
		section.bottom_margin = Inches(bottom_margin_inch)

		# Section type - continuous
		section.start_type = 2  # WD_SECTION.CONTINUOUS

	# Apply font to all paragraphs
	for paragraph in doc.paragraphs:
		for run in paragraph.runs:
			run.font.name = font_name
			run.font.size = Pt(font_size)

	# Add table borders, apply font, set full width and equal columns
	for table in doc.tables:
		# Set table width to 100% using twips
		tbl = table._element
		tblPr = tbl.tblPr
		if tblPr is None:
			tblPr = OxmlElement("w:tblPr")
			tbl.insert(0, tblPr)

		# Remove table cell spacing
		tblCellSpacing = tblPr.find(qn("w:tblCellSpacing"))
		if tblCellSpacing is not None:
			tblPr.remove(tblCellSpacing)

		# Set table width in twips (absolute measurement)
		tblW = tblPr.find(qn("w:tblW"))
		if tblW is None:
			tblW = OxmlElement("w:tblW")
			tblPr.append(tblW)
		tblW.set(qn("w:w"), str(available_width_twips))
		tblW.set(qn("w:type"), "dxa")  # Use absolute width in twips

		# Set table alignment to left (to fill from left margin)
		tblJc = tblPr.find(qn("w:jc"))
		if tblJc is None:
			tblJc = OxmlElement("w:jc")
			tblPr.append(tblJc)
		tblJc.set(qn("w:val"), "left")

		# Set table layout to fixed
		tblLayout = tblPr.find(qn("w:tblLayout"))
		if tblLayout is None:
			tblLayout = OxmlElement("w:tblLayout")
			tblPr.append(tblLayout)
		tblLayout.set(qn("w:type"), "fixed")

		# Remove table indentation
		tblInd = tblPr.find(qn("w:tblInd"))
		if tblInd is None:
			tblInd = OxmlElement("w:tblInd")
			tblPr.append(tblInd)
		tblInd.set(qn("w:w"), "0")
		tblInd.set(qn("w:type"), "dxa")

		# Remove table borders at table level and set at cell level only
		tblBorders = tblPr.find(qn("w:tblBorders"))
		if tblBorders is not None:
			tblPr.remove(tblBorders)

		# Get number of columns from first row
		num_cols = 0
		if table.rows:
			num_cols = len(table.rows[0].cells)

		# Calculate equal column width in twips
		if num_cols > 0:
			equal_col_width = int(available_width_twips / num_cols)

		# Update grid columns
		tblGrid = tbl.find(qn("w:tblGrid"))
		if tblGrid is not None:
			# Clear existing grid columns
			for gridCol in list(tblGrid):
				tblGrid.remove(gridCol)
			# Add new grid columns with equal width
			for _ in range(num_cols):
				gridCol = OxmlElement("w:gridCol")
				gridCol.set(qn("w:w"), str(equal_col_width))
				tblGrid.append(gridCol)

		for row in table.rows:
			for cell in row.cells:
				# Apply font to cell text
				for paragraph in cell.paragraphs:
					for run in paragraph.runs:
						run.font.name = font_name
						run.font.size = Pt(font_size)

				# Set equal cell width in twips
				tc = cell._tc
				tcPr = tc.get_or_add_tcPr()

				# Add cell width property
				tcW = tcPr.find(qn("w:tcW"))
				if tcW is None:
					tcW = OxmlElement("w:tcW")
					tcPr.append(tcW)
				if num_cols > 0:
					tcW.set(qn("w:w"), str(equal_col_width))
					tcW.set(qn("w:type"), "dxa")  # Use absolute width

				# Add cell margins for better spacing
				tcMar = tcPr.find(qn("w:tcMar"))
				if tcMar is None:
					tcMar = OxmlElement("w:tcMar")
					tcPr.append(tcMar)

				# Set cell padding (reduced for more space)
				for margin_side in ["left", "right"]:
					margin = OxmlElement(f"w:{margin_side}")
					margin.set(qn("w:w"), "50")  # Reduced from 100
					margin.set(qn("w:type"), "dxa")
					tcMar.append(margin)
				for margin_side in ["top", "bottom"]:
					margin = OxmlElement(f"w:{margin_side}")
					margin.set(qn("w:w"), "50")
					margin.set(qn("w:type"), "dxa")
					tcMar.append(margin)

				# Add borders
				tcBorders = tcPr.first_child_found_in("w:tcBorders")
				if tcBorders is None:
					tcBorders = OxmlElement("w:tcBorders")
					tcPr.append(tcBorders)
				else:
					# Clear existing borders
					for border in list(tcBorders):
						tcBorders.remove(border)

				for border_name in ["top", "left", "bottom", "right"]:
					border = OxmlElement(f"w:{border_name}")
					border.set(qn("w:val"), "single")
					border.set(qn("w:sz"), "4")
					border.set(qn("w:space"), "0")
					border.set(qn("w:color"), "000000")
					tcBorders.append(border)

	# Return updated DOCX bytes
	new_stream = BytesIO()
	doc.save(new_stream)
	new_stream.seek(0)
	return new_stream.read()


def send_file(bytes_data, filename, filetype="download"):
	frappe.local.response.filename = filename
	frappe.local.response.filecontent = bytes_data
	frappe.local.response.type = filetype


@frappe.whitelist()
def generate_pdf_template(template_path_or_print_format, filename=None, **kwargs):
	"""
	Generates PDF by default, DOCX if kwargs['type'] == 'docx'
	Formatting parameters only apply to DOCX generation
	"""
	try:
		# Render template
		if (
			isinstance(template_path_or_print_format, str)
			and "templates/pages" in template_path_or_print_format
		):
			template_html = frappe.get_template(template_path_or_print_format).render(kwargs)
		else:
			if frappe.db.exists("Print Format", template_path_or_print_format):
				print_format_template = frappe.get_doc("Print Format", template_path_or_print_format).html
				template_html = frappe.render_template(print_format_template, kwargs)
			else:
				frappe.throw(f"Print Format not found: {template_path_or_print_format}")

		# Generate filename with date
		today = nowdate()
		formatted_today = formatdate(today, "dd-MM-yyyy")
		base_filename = (filename or template_path_or_print_format) + f"_{formatted_today}"

		# Conditional PDF or DOCX
		if kwargs.get("type") == "docx":
			pdf_bytes = get_pdf(template_html)
			docx_bytes = pdf_bytes_to_docx_bytes(pdf_bytes)
			send_file(docx_bytes, f"{base_filename}.docx")
		else:
			# PDF generation - no special formatting applied
			pdf_bytes = get_pdf(template_html)
			send_file(pdf_bytes, f"{base_filename}.pdf")

	except Exception as e:
		frappe.log_error(f"Error generating document: {str(e)}")
		frappe.throw(f"Error generating document: {str(e)}")
